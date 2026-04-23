"""
build_templates.py v3 — căn chỉnh chính xác theo XML mẫu gốc MobiFone Lâm Đồng.
Chạy: python build_templates.py
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = "word_templates"
os.makedirs(OUT, exist_ok=True)

FONT = "Times New Roman"

# ── DXA constants từ XML mẫu gốc ─────────────────────────────────────────────
# Header table
HDR_LEFT_TT  = 3330
HDR_RIGHT_TT = 6379
HDR_LEFT_BC  = 6777
HDR_RIGHT_BC = 6777
GDN_LEFT     = 5098
GDN_RIGHT    = 3918

# Ký tên tờ trình 3 cột
KY_TT_C0 = 2410
KY_TT_C1 = 2694
KY_TT_C2 = 4282

# Báo cáo KQCV — column widths từ mẫu
BC_COLS = [895, 3623, 2497, 2880, 1800, 1859]
KY_BC   = 6777

# Font sizes (half-points: 26=13pt, 28=14pt, 32=16pt)
SZ_TITLE = 32
SZ_HEAD2 = 30
SZ_BODY  = 26
SZ_HDR   = 26
SZ_SMALL = 22

# Line spacing
LS_BODY  = 276
LS_TIGHT = 264


# ── helpers ───────────────────────────────────────────────────────────────────

def set_margins(doc, top=2.5, bot=2.5, left=3.0, right=2.0):
    for s in doc.sections:
        s.top_margin    = Cm(top)
        s.bottom_margin = Cm(bot)
        s.left_margin   = Cm(left)
        s.right_margin  = Cm(right)

def _font(run, sz):
    """Set font name + sz + szCs (cần cho tiếng Việt)."""
    rPr = run._r.get_or_add_rPr()
    # rFonts
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rPr.insert(0, rf)
    for attr in ['w:ascii','w:hAnsi','w:cs','w:eastAsia']:
        rf.set(qn(attr), FONT)
    # sz + szCs
    for tag in ['w:sz','w:szCs']:
        el = rPr.find(qn(tag))
        if el is not None:
            rPr.remove(el)
        el = OxmlElement(tag)
        el.set(qn('w:val'), str(sz))
        rPr.append(el)

def _run(p, text, bold=False, italic=False, sz=SZ_BODY):
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    _font(r, sz)
    return r

def _pPr(p, jc='both', line=LS_BODY, before=120, after=120,
         first_line=None, start=None, hanging=None, end=0):
    """Áp XML formatting trực tiếp vào paragraph."""
    pp = p._p.get_or_add_pPr()
    # jc
    jc_el = pp.find(qn('w:jc'))
    if jc_el is None:
        jc_el = OxmlElement('w:jc'); pp.append(jc_el)
    jc_el.set(qn('w:val'), jc)
    # spacing
    sp = pp.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing'); pp.append(sp)
    sp.set(qn('w:line'),     str(line))
    sp.set(qn('w:lineRule'), 'auto')
    if before is not None: sp.set(qn('w:before'), str(before))
    if after  is not None: sp.set(qn('w:after'),  str(after))
    # indent
    if any(v is not None for v in [first_line, start, hanging]):
        ind = pp.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind'); pp.append(ind)
        if first_line is not None: ind.set(qn('w:firstLine'), str(first_line))
        if start      is not None: ind.set(qn('w:start'),     str(start))
        if hanging    is not None: ind.set(qn('w:hanging'),   str(hanging))
        ind.set(qn('w:end'), str(end))

def _para(doc, text='', bold=False, italic=False, sz=SZ_BODY,
          jc='left', line=LS_BODY, before=120, after=120,
          first_line=None, start=None, hanging=None):
    p = doc.add_paragraph()
    _pPr(p, jc=jc, line=line, before=before, after=after,
         first_line=first_line, start=start, hanging=hanging)
    if text: _run(p, text, bold=bold, italic=italic, sz=sz)
    return p

def _no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for s in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'none')
        tcB.append(b)
    tcPr.append(tcB)

def _tbl_no_border(tbl):
    tblPr = tbl._tbl.tblPr
    tblB = OxmlElement('w:tblBorders')
    for s in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'none')
        tblB.append(b)
    tblPr.append(tblB)

def _tbl_border(tbl):
    tblPr = tbl._tbl.tblPr
    tblB = OxmlElement('w:tblBorders')
    for s in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        tblB.append(b)
    tblPr.append(tblB)

def _tbl_w(tbl, w, t='dxa'):
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(w))
    tblW.set(qn('w:type'), t)

def _cell_w(cell, w):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(w))
    tcW.set(qn('w:type'), 'dxa')

def _cp(cell, text='', bold=False, italic=False, sz=SZ_BODY,
        jc='center', line=LS_BODY, before=60, after=60, new=False):
    """Viết vào ô bảng."""
    p = cell.add_paragraph() if new else cell.paragraphs[0]
    _pPr(p, jc=jc, line=line, before=before, after=after)
    if text: _run(p, text, bold=bold, italic=italic, sz=sz)
    return p

def _cpa(cell, text='', bold=False, italic=False, sz=SZ_BODY,
         jc='center', line=LS_BODY, before=60, after=60):
    return _cp(cell, text, bold, italic, sz, jc, line, before, after, new=True)

def quoc_hieu(doc, left_rows, right_rows, lw, rw):
    tbl = doc.add_table(rows=1, cols=2)
    _tbl_no_border(tbl)
    _tbl_w(tbl, lw + rw, 'dxa')
    row = tbl.rows[0]
    _cell_w(row.cells[0], lw)
    _cell_w(row.cells[1], rw)
    for cell, rows in [(row.cells[0], left_rows), (row.cells[1], right_rows)]:
        _no_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        first = True
        for item in rows:
            t    = item[0]
            b    = item[1] if len(item)>1 else False
            sz   = item[2] if len(item)>2 else SZ_HDR
            ital = item[3] if len(item)>3 else False
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            _pPr(p, jc='center', line=LS_BODY, before=0, after=30)
            if t: _run(p, t, bold=b, italic=ital, sz=sz)
    _para(doc, '', before=0, after=60, line=LS_BODY)


# ── 1. TỜ TRÌNH ──────────────────────────────────────────────────────────────

def make_to_trinh():
    doc = Document()
    set_margins(doc)

    quoc_hieu(doc,
        [('MOBIFONE LÂM ĐỒNG', True, SZ_HDR),
         ('TT KD GIẢI PHÁP SỐ', True, SZ_HDR),
         ('Số: {{so_to_trinh}}/TTr-TTKDGPS', False, SZ_HDR)],
        [('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM', True, SZ_HDR),
         ('Độc lập – Tự do – Hạnh phúc', False, SZ_HDR),
         ('', False, SZ_HDR),
         ('Lâm Đồng, {{ngay_to_trinh}}', False, SZ_HDR, True)],
        HDR_LEFT_TT, HDR_RIGHT_TT)

    _para(doc, 'TỜ TRÌNH', bold=True, sz=SZ_HEAD2,
          jc='center', before=60, after=60, line=LS_BODY)

    p = _para(doc, '', sz=SZ_BODY, jc='center', before=0, after=120, line=LS_BODY)
    _run(p, 'Về việc: ', bold=True, sz=SZ_BODY)
    _run(p, 'Đề xuất chi phí tiếp khách tại MobiFone Lâm Đồng tháng {{thang_tt}}',
         bold=True, sz=SZ_BODY)

    # Kính trình — firstLine=180, start=180, line=360 (double space) theo mẫu
    _para(doc, 'Kính trình: Ông {{giam_doc}} – Giám đốc MobiFone Lâm Đồng',
          sz=SZ_BODY, jc='center', line=360, before=120, after=120,
          first_line=180, start=180)

    # Ô ý kiến — 9000 dxa theo mẫu
    tyk = doc.add_table(rows=1, cols=1)
    tyk.style = 'Table Grid'
    _tbl_border(tyk)
    _tbl_w(tyk, 9000, 'dxa')
    c = tyk.rows[0].cells[0]
    _cell_w(c, 9000)
    _cp(c, 'Ý kiến chỉ đạo của Lãnh đạo:', bold=True, sz=SZ_BODY,
        jc='left', before=80, after=500)

    _para(doc, '', before=0, after=60, line=LS_BODY)

    # Body — jc=both firstLine=567, line=264, before=120 after=120
    def B(text, after=120, fl=567, st=0): 
        _para(doc, text, sz=SZ_BODY, jc='both', line=LS_TIGHT,
              before=120, after=after, first_line=fl, start=st)

    B('Căn cứ Quyết định số {{quyet_dinh_cp}} về việc giao kế hoạch chi phí lần 1 năm {{nam_tt_so}};')
    B('Căn cứ nhu cầu thực tế tại đơn vị.')
    B('Trung tâm Kinh doanh Giải pháp số kính trình Lãnh đạo về việc duyệt chi phí tiếp khách tại MobiFone Lâm Đồng, cụ thể như sau:')

    def bh(label):  # bullet heading: hanging=425 start=709
        p = _para(doc, '', sz=SZ_BODY, jc='both', line=LS_TIGHT,
                  before=120, after=120, hanging=425, start=709)
        _run(p, '- ', bold=True, sz=SZ_BODY)
        _run(p, label, bold=True, sz=SZ_BODY)

    def bi(text, after=0):  # bullet item: firstLine=284 start=142
        _para(doc, text, sz=SZ_BODY, jc='both', line=LS_TIGHT,
              before=120, after=after, first_line=284, start=142)

    def bi2(text, after=0): # bullet item thụt: firstLine=426
        _para(doc, text, sz=SZ_BODY, jc='both', line=LS_TIGHT,
              before=120, after=after, first_line=426, start=0)

    bh('Mục tiêu, lý do:')
    bi('- Nhằm đảm bảo sản xuất kinh doanh tại MobiFone Lâm Đồng.')
    bi2('- {{ly_do}}', after=120)

    bh('Nội dung thực hiện:')
    bi('- Mời cơm thân mật cùng với các Doanh nghiệp trên địa bàn.')
    bi('- Trao đổi, giải đáp thắc mắc của doanh nghiệp đối với các sản phẩm của MobiFone.')
    bi2('- Thành phần tham dự: Lãnh đạo, chuyên viên liên quan, đại diện các Doanh nghiệp.', after=120)

    p = _para(doc, '', sz=SZ_BODY, jc='both', line=LS_BODY,
              before=120, after=120, first_line=284, start=0)
    _run(p, '- ', sz=SZ_BODY)
    _run(p, 'Chi phí dự kiến: ', bold=True, sz=SZ_BODY)
    _run(p, '{{truoc_vat}}đ (chưa bao gồm VAT). Bằng chữ: {{tien_bang_chu}}.', sz=SZ_BODY)

    p = _para(doc, '', sz=SZ_BODY, jc='both', line=LS_BODY,
              before=120, after=120, first_line=284, start=0)
    _run(p, '- ', sz=SZ_BODY)
    _run(p, 'Nguồn chi phí:', bold=True, sz=SZ_BODY)

    bi('- Nguồn chi phí: Quyết định số {{quyet_dinh_cp}}', after=0)
    bi('- KMCP: {{ma_kmcp}} – MOBIFONE - VT - Di động - CPQLDN - CP Tiếp khách, khánh tiết - Quản lý hành chính', after=120)

    # Thời gian — hanging=142, start=426
    p = _para(doc, '', sz=SZ_BODY, jc='left', line=LS_TIGHT,
              before=120, after=120, hanging=142, start=426)
    _run(p, '- ', sz=SZ_BODY)
    _run(p, 'Thời gian thực hiện: ', bold=True, sz=SZ_BODY)
    _run(p, 'Tháng {{thang_tt}}', sz=SZ_BODY)

    B('Kính trình Lãnh đạo xem xét và phê duyệt cho Trung tâm Kinh doanh Giải pháp số được chủ trì bố trí kế hoạch cụ thể trên địa bàn nhằm phù hợp với thực tế./.', after=240)

    # Ký 3 cột — 2410, 2694, 4282
    tky = doc.add_table(rows=1, cols=3)
    _tbl_no_border(tky)
    _tbl_w(tky, KY_TT_C0+KY_TT_C1+KY_TT_C2, 'dxa')
    row = tky.rows[0]
    for i, w in enumerate([KY_TT_C0, KY_TT_C1, KY_TT_C2]):
        _cell_w(row.cells[i], w)
        _no_border(row.cells[i])

    _cp(row.cells[0], 'Nơi nhận:', bold=True, italic=True, sz=SZ_SMALL, jc='left', before=60, after=0)
    _cpa(row.cells[0], '- Như trên;', italic=True, sz=SZ_SMALL, jc='left', before=0, after=0)
    _cpa(row.cells[0], '- Lưu P.TH.', italic=True, sz=SZ_SMALL, jc='left', before=0, after=0)

    _cp(row.cells[1], '', sz=SZ_HDR, jc='center', before=60, after=0)

    _cp(row.cells[2], 'PHỤ TRÁCH TRUNG TÂM', bold=True, sz=SZ_HDR, jc='center', before=60, after=0)
    for _ in range(4):
        _cpa(row.cells[2], '', sz=SZ_HDR, jc='center', before=0, after=0)
    _cpa(row.cells[2], '{{lanh_dao}}', bold=True, sz=SZ_HDR, jc='center', before=0, after=0)

    doc.save(f'{OUT}/to_trinh.docx')
    print('✓ to_trinh.docx')


# ── 2. GIẤY ĐỀ NGHỊ TIẾP KHÁCH ──────────────────────────────────────────────

def make_giay_de_nghi():
    doc = Document()
    set_margins(doc)

    # Header — cell0=5098 center, cell1=3918 RIGHT (theo mẫu gốc)
    tbl_h = doc.add_table(rows=1, cols=2)
    _tbl_no_border(tbl_h)
    _tbl_w(tbl_h, GDN_LEFT+GDN_RIGHT, 'dxa')
    row = tbl_h.rows[0]
    _cell_w(row.cells[0], GDN_LEFT)
    _cell_w(row.cells[1], GDN_RIGHT)
    for c in row.cells:
        _no_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Cell trái
    _cp(row.cells[0], 'TCT VIỄN THÔNG MOBIFONE', sz=SZ_HDR, jc='center', before=0, after=30)
    _cpa(row.cells[0], 'MOBIFONE LÂM ĐỒNG', bold=True, sz=SZ_HDR, jc='center', before=0, after=30)

    # Cell phải — align RIGHT theo mẫu gốc
    _cp(row.cells[1], 'MẪU SỐ: 08-TT', sz=SZ_HDR, jc='right', before=0, after=0)

    _para(doc, '', before=0, after=60, line=LS_BODY)

    # Tiêu đề sz=32
    _para(doc, 'GIẤY ĐỀ NGHỊ TIẾP KHÁCH', bold=True, sz=SZ_TITLE,
          jc='center', before=60, after=60, line=LS_BODY)

    # Ngày + Kính gửi sz=28
    _para(doc, '{{ngay_tiep_khach}}', italic=True, sz=28,
          jc='center', before=0, after=80, line=LS_BODY)

    p = _para(doc, '', sz=28, jc='center', before=0, after=80, line=LS_BODY)
    _run(p, 'Kính gửi: ', bold=True, sz=28)
    _run(p, 'Lãnh đạo MobiFone Lâm Đồng', bold=True, sz=28)

    # Body sz=28, jc=both, firstLine=567 (theo mẫu gốc)
    def GB(text, fl=567, after=120, bold=False):
        _para(doc, text, bold=bold, sz=28, jc='both', line=LS_BODY,
              before=120, after=after, first_line=fl, start=0)

    GB('Đơn vị đề nghị: {{don_vi}} – MobiFone Lâm Đồng')
    GB('Lý do: {{ly_do}}', fl=0)
    GB('Người chủ trì: {{ho_ten}}')
    GB('Thành phần và số người tham gia')
    GB('+ Về phía MBF Lâm Đồng: {{sl_ld}} Lãnh đạo và {{sl_cv}} Chuyên viên')
    GB('+ Về phía khách mời: {{sl_khach}} người')
    GB('Khách mời: {{khach_moi}}')
    GB('Dự trù số tiền chi tiếp khách: {{truoc_vat}}đ (chưa VAT). Bằng chữ: {{tien_bang_chu}}.', after=240)

    # Ký 2 cột — 4798, 4822
    tky = doc.add_table(rows=2, cols=2)
    _tbl_no_border(tky)
    _tbl_w(tky, 4798+4822, 'dxa')
    for row in tky.rows:
        _cell_w(row.cells[0], 4798)
        _cell_w(row.cells[1], 4822)
        for c in row.cells: _no_border(c)

    _cp(tky.rows[0].cells[0], 'Thủ trưởng đơn vị', bold=True, sz=28, jc='center', before=60, after=0)
    _cpa(tky.rows[0].cells[0], '(Ký, họ tên)', sz=28, jc='center', before=0, after=0)
    _cp(tky.rows[0].cells[1], 'Người đề nghị', bold=True, sz=28, jc='center', before=60, after=0)
    _cpa(tky.rows[0].cells[1], '(Ký, họ tên)', sz=28, jc='center', before=0, after=0)

    for ci, name in [(0,'{{lanh_dao}}'),(1,'{{ho_ten}}')]:
        _cp(tky.rows[1].cells[ci], '', sz=28, jc='center', before=0, after=0)
        for _ in range(3):
            _cpa(tky.rows[1].cells[ci], '', sz=28, jc='center', before=0, after=0)
        _cpa(tky.rows[1].cells[ci], name, bold=True, sz=28, jc='center', before=0, after=0)

    doc.save(f'{OUT}/giay_de_nghi.docx')
    print('✓ giay_de_nghi.docx')


# ── 3. BẢNG KÊ ───────────────────────────────────────────────────────────────

def make_bang_ke():
    doc = Document()
    set_margins(doc, top=2, bot=2, left=2.5, right=2.0)

    _para(doc, 'BẢNG KÊ THANH TOÁN', bold=True, sz=SZ_TITLE,
          jc='center', before=0, after=0, line=LS_BODY)
    _para(doc, 'CHI PHÍ NĂM {{nam_bc_so}}', bold=True, sz=SZ_TITLE,
          jc='center', before=0, after=120, line=LS_BODY)
    _para(doc, 'Họ và tên:\t\t{{ho_ten}}', sz=SZ_BODY, jc='left',
          before=60, after=60, line=LS_BODY)
    _para(doc, 'Đơn vị:\t\tMobiFone Lâm Đồng', sz=SZ_BODY, jc='left',
          before=0, after=120, line=LS_BODY)

    # Cột bảng kê 11 cột
    BK_W = [512, 910, 1820, 853, 2274, 853, 740, 568, 1080, 853, 1080]
    hdrs = ['STT','Số chứng từ','Nội dung CV','Mã KM','Tên KM',
            'Tài khoản','Nghiệp vụ','Mã SPDV','Trước VAT','VAT','Sau VAT']

    tbl = doc.add_table(rows=3, cols=11)
    _tbl_border(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl_w(tbl, sum(BK_W), 'dxa')
    for j, w in enumerate(BK_W):
        for row in tbl.rows: _cell_w(row.cells[j], w)

    for j, h in enumerate(hdrs):
        _cp(tbl.rows[0].cells[j], h, bold=True, sz=SZ_SMALL, jc='center', before=60, after=60)

    data = ['1','{{so_hd}}',
            'CHI TIEP KHACH T{{thang_tt}} - {{ma_kmcp}}',
            '{{ma_kmcp}}',
            'MOBIFONE - VT - Di động - CPQLDN - CP Tiếp khách, khánh tiết - Quản lý hành chính',
            '{{tk_kt}}','{{nghiep_vu}}','{{ma_spdv}}',
            '{{truoc_vat}}','{{tien_vat}}','{{sau_vat}}']
    for j, v in enumerate(data):
        _cp(tbl.rows[1].cells[j], v, sz=SZ_SMALL,
            jc='center' if j in {0,1,3,5,6,7,8,9,10} else 'both',
            before=60, after=60)

    tbl.rows[2].cells[0].merge(tbl.rows[2].cells[7])
    _cp(tbl.rows[2].cells[0], 'Tổng cộng', bold=True, sz=SZ_BODY, jc='center', before=60, after=60)
    for j, k in [(8,'{{truoc_vat}}'),(9,'{{tien_vat}}'),(10,'{{sau_vat}}')]:
        _cp(tbl.rows[2].cells[j], k, bold=True, sz=SZ_BODY, jc='center', before=60, after=60)

    _para(doc, '', before=0, after=60, line=LS_BODY)
    _para(doc, '{{ngay_bao_cao}}', italic=True, sz=SZ_BODY,
          jc='right', before=0, after=60, line=LS_BODY)

    tky = doc.add_table(rows=2, cols=2)
    _tbl_no_border(tky)
    _tbl_w(tky, 9354, 'dxa')
    for row in tky.rows:
        _cell_w(row.cells[0], 4677); _cell_w(row.cells[1], 4677)
        for c in row.cells: _no_border(c)

    _cp(tky.rows[0].cells[0], 'PHỤ TRÁCH CHI PHÍ ĐƠN VỊ', bold=True, sz=SZ_BODY, jc='center', before=60, after=0)
    _cpa(tky.rows[0].cells[0], '(Ghi rõ họ, tên)', sz=SZ_BODY, jc='center', before=0, after=0)
    _cp(tky.rows[0].cells[1], 'NGƯỜI ĐỀ NGHỊ', bold=True, sz=SZ_BODY, jc='center', before=60, after=0)
    _cpa(tky.rows[0].cells[1], '(Ghi rõ họ, tên)', sz=SZ_BODY, jc='center', before=0, after=0)

    for ci, nm in [(0,'{{phu_trach_cp}}'),(1,'{{ho_ten}}')]:
        _cp(tky.rows[1].cells[ci], '', sz=SZ_BODY, jc='center', before=0, after=0)
        for _ in range(3): _cpa(tky.rows[1].cells[ci], '', sz=SZ_BODY, jc='center', before=0, after=0)
        _cpa(tky.rows[1].cells[ci], nm, bold=True, sz=SZ_BODY, jc='center', before=0, after=0)

    # Trang 2: Bảng kê khai
    doc.add_page_break()
    _para(doc, 'TỔNG CÔNG TY VIỄN THÔNG MOBIFONE', bold=True, sz=SZ_HDR, jc='center', before=0, after=0)
    _para(doc, 'Đơn vị: MOBIFONE LÂM ĐỒNG', sz=SZ_HDR, jc='center', before=0, after=120)
    _para(doc, 'BẢNG KÊ KHAI CHI PHÍ THƯỜNG XUYÊN THANH TOÁN', bold=True, sz=SZ_BODY, jc='center', before=0, after=0)
    _para(doc, 'THÁNG {{thang_tk_so}} NĂM {{nam_tk_so}}', bold=True, sz=SZ_BODY, jc='center', before=0, after=120)

    BKK_W = [512, 2500, 1420, 1140, 1420, 1000, 1420]
    hdrs2 = ['STT','Loại chi phí','Tháng\nthanh toán','Hóa đơn','Trước VAT','VAT','Sau VAT']

    tbl3 = doc.add_table(rows=3, cols=7)
    _tbl_border(tbl3)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl_w(tbl3, sum(BKK_W), 'dxa')
    for j, w in enumerate(BKK_W):
        for row in tbl3.rows: _cell_w(row.cells[j], w)

    for j, h in enumerate(hdrs2):
        _cp(tbl3.rows[0].cells[j], h, bold=True, sz=SZ_BODY, jc='center', before=60, after=60)

    data2 = ['1','Chi phí tiếp khách T{{thang_tk_so}}.{{nam_tk_so}}',
             '{{thang_tt}}','{{so_hd}}','{{truoc_vat}}','{{tien_vat}}','{{sau_vat}}']
    for j, v in enumerate(data2):
        _cp(tbl3.rows[1].cells[j], v, sz=SZ_BODY,
            jc='center' if j in {0,2,3,4,5,6} else 'left', before=60, after=60)

    tbl3.rows[2].cells[0].merge(tbl3.rows[2].cells[3])
    _cp(tbl3.rows[2].cells[0], 'Tổng cộng', bold=True, sz=SZ_BODY, jc='center', before=60, after=60)
    for j, k in [(4,'{{truoc_vat}}'),(5,'{{tien_vat}}'),(6,'{{sau_vat}}')]:
        _cp(tbl3.rows[2].cells[j], k, bold=True, sz=SZ_BODY, jc='center', before=60, after=60)

    _para(doc, '', before=60, after=60)
    _para(doc, '\t\t\t\tNgười lập bảng', sz=SZ_BODY, jc='left', before=0, after=0)
    for _ in range(3): _para(doc, '', before=0, after=0)
    _para(doc, '\t\t\t\t{{ho_ten}}', bold=True, sz=SZ_BODY, jc='left', before=0, after=0)

    doc.save(f'{OUT}/bang_ke.docx')
    print('✓ bang_ke.docx')


# ── 4. BÁO CÁO KẾT QUẢ CÔNG VIỆC ────────────────────────────────────────────

def make_bao_cao_kqcv():
    doc = Document()
    set_margins(doc)

    quoc_hieu(doc,
        [('TỔNG CÔNG TY VIỄN THÔNG MOBIFONE', False, SZ_HDR),
         ('MOBIFONE LÂM ĐỒNG', True, SZ_HDR)],
        [('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM', True, SZ_HDR),
         ('Độc lập – Tự do – Hạnh phúc', False, SZ_HDR),
         ('', False, SZ_HDR),
         ('Lâm Đồng, {{ngay_bao_cao}}', False, SZ_HDR, True)],
        HDR_LEFT_BC, HDR_RIGHT_BC)

    _para(doc, 'BÁO CÁO KẾT QUẢ CÔNG VIỆC', bold=True, sz=SZ_TITLE,
          jc='center', before=60, after=120)

    # Body sz=26, jc=left (theo mẫu gốc — không justify)
    for txt in ['Họ và tên: {{ho_ten}}',
                'Tổ/ Bộ phận công tác: {{don_vi}}',
                'Nội dung công việc: Tiếp khách',
                'Thành phần: Lãnh đạo Chi nhánh, chuyên viên và khách mời (chi tiết theo tờ trình)']:
        _para(doc, txt, sz=SZ_BODY, jc='left', before=0, after=0)
    _para(doc, '', before=0, after=60)

    # Bảng KQCV — widths 895,3623,2497,2880,1800,1859 từ mẫu
    tbl = doc.add_table(rows=2, cols=6)
    _tbl_border(tbl)
    _tbl_w(tbl, 0, 'auto')
    for j, w in enumerate(BC_COLS):
        for row in tbl.rows: _cell_w(row.cells[j], w)

    hdrs = ['STT','NỘI DUNG CÔNG VIỆC','THỜI GIAN LÀM VIỆC',
            'KẾT QUẢ ĐẠT ĐƯỢC','CHƯA ĐẠT ĐƯỢC','HƯỚNG GIẢI QUYẾT']
    for j, h in enumerate(hdrs):
        _cp(tbl.rows[0].cells[j], h, bold=True, sz=SZ_BODY, jc='center', before=60, after=60)

    data = ['01','{{ly_do}}',
            '{{ngay_tk_so}}/{{thang_tk_so}}/{{nam_tk_so}}',
            '{{ket_qua}}','','']
    for j, v in enumerate(data):
        _cp(tbl.rows[1].cells[j], v, sz=SZ_BODY,
            jc='center' if j in {0,2} else 'both',
            line=276, before=60, after=60)

    _para(doc, '', before=60, after=120)

    # Ký — 6777 x 2
    tky = doc.add_table(rows=2, cols=2)
    _tbl_no_border(tky)
    _tbl_w(tky, KY_BC*2, 'dxa')
    for row in tky.rows:
        _cell_w(row.cells[0], KY_BC); _cell_w(row.cells[1], KY_BC)
        for c in row.cells: _no_border(c)

    _cp(tky.rows[0].cells[0], '', sz=SZ_BODY, jc='center', before=60, after=0)
    _cp(tky.rows[0].cells[1], 'NGƯỜI BÁO CÁO', bold=True, sz=SZ_BODY, jc='center', before=60, after=0)
    _cp(tky.rows[1].cells[0], '', sz=SZ_BODY, jc='center', before=0, after=0)
    _cp(tky.rows[1].cells[1], '', sz=SZ_BODY, jc='center', before=0, after=0)
    for _ in range(3): _cpa(tky.rows[1].cells[1], '', sz=SZ_BODY, jc='center', before=0, after=0)
    _cpa(tky.rows[1].cells[1], '{{ho_ten}}', bold=True, sz=SZ_BODY, jc='center', before=0, after=0)

    doc.save(f'{OUT}/bao_cao_kqcv.docx')
    print('✓ bao_cao_kqcv.docx')


if __name__ == '__main__':
    make_to_trinh()
    make_giay_de_nghi()
    make_bang_ke()
    make_bao_cao_kqcv()
    print('\nDone — 4 templates pixel-perfect trong word_templates/')
